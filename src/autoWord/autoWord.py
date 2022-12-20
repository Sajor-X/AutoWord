# Author: @Sajor
# License: MIT
import uuid

import click as click
from docx import Document


def generate_uuid() -> str:
    """
    Generate a UUID for the session -- Internal use only

    :return: a random UUID
    :rtype: :obj:`str`
    """
    uid = str(uuid.uuid4())
    return uid


class Debugger:
    def __init__(self, debug: bool = False):
        if debug:
            print("Debugger enabled on AutoWord")
        self.debug = debug

    def set_debug(self, debug: bool):
        self.debug = debug

    def log(self, message: str, end: str = "\n"):
        if self.debug:
            print(message, end=end)


class AutoWord:

    def __init__(self, debug=False):
        self.debugger = Debugger(debug)

    def generate_file(self, title, body, output_path):
        """
        auto generator docx file
        :param title: 标题
        :param body: 正文
        :param output_path: 输出路径
        """
        document = Document()  # 生成doc对象
        document.add_heading(title, 0)  # 添加标题
        document.add_paragraph(body)  # 写入正文
        document.save(output_path)  # 保存文件
        self.debugger.log("The generated completed.")


@click.command()
@click.option('--title', '-t', prompt="Article Title", default="Hello", help='Article title')
@click.option('--body', '-b', help='Article body')
@click.option('--path', '-p', prompt="Article Path", default="Hello.docx", help='Article path')
def cli(title, body, path):
    c = AutoWord()
    c.generate_file(title, body, path)


if __name__ == '__main__':
    cli()
